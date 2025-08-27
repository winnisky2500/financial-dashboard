#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Report Agent (LOCKED PROMPT VERSION)
- 前端仅提供：公司、起止年季、语言、特殊要求、报告模板（ID 或模板 JSON）
- 提示词/写作风格/准确性规则/图表规范全部在后端固定，前端无法覆盖
- 从 financial_metrics 拉数；从 policy_news/可选外网检索拿政策上下文
- 调用大模型生成 Markdown（含 ECharts 占位）；可选导出 DOCX/PDF 至 Supabase Storage

环境变量（仅后端）：
  SUPABASE_URL=...
  SUPABASE_SERVICE_ROLE_KEY=...
  REPORTS_BUCKET=reports
  REPORT_AGENT_TOKEN=dev-secret-01

  OPENAI_API_KEY=...                 # 本地放 .env.local；线上放部署平台 Secrets
  OPENAI_BASE_URL=https://api.openai.com/v1
  OPENAI_MODEL=gpt-4o-mini

  # 可选外部检索（二选一；不配则仅用 policy_news）
  BING_SUBSCRIPTION_KEY=...
  BING_ENDPOINT=https://api.bing.microsoft.com/v7.0/search
  或
  GOOGLE_API_KEY=...
  GOOGLE_CSE_ID=...

  EXPORT_ENABLED=1                   # 1=导出并上传PDF/DOCX；0=仅返回Markdown
"""

import os, io, uuid, json, datetime as dt
from typing import Optional, List, Dict, Any
import requests
import pandas as pd

from fastapi import FastAPI, HTTPException, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

#from supabase import create_client, Client
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

from supabase import create_client, Client
import logging, traceback
import re


logger = logging.getLogger("report_agent")
logger.setLevel(logging.INFO)

# ✅ 支持两套环境变量名
SUPABASE_URL = os.getenv("SUPABASE_URL") or os.getenv("VITE_SUPABASE_URL", "")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY") or os.getenv("VITE_SUPABASE_SERVICE_ROLE_KEY", "")

REPORTS_BUCKET = os.getenv("REPORTS_BUCKET", "reports")
REPORT_AGENT_TOKEN = os.getenv("REPORT_AGENT_TOKEN", "dev-secret-01")
EXPORT_ENABLED = os.getenv("EXPORT_ENABLED", "1") == "1"

# LLM（OpenAI 兼容）
from openai import OpenAI
import httpx
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
USE_LLM = bool(OPENAI_API_KEY)

# 外部搜索（可选）
BING_SUBSCRIPTION_KEY = os.getenv("BING_SUBSCRIPTION_KEY", "")
BING_ENDPOINT = os.getenv("BING_ENDPOINT", "https://api.bing.microsoft.com/v7.0/search")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
GOOGLE_CSE_ID = os.getenv("GOOGLE_CSE_ID", "")

# === Token 预算与瘦身工具 ===
def guess_context_limit(model: str) -> int:
    m = (model or "").lower()
    if "4o" in m or "o3" in m:    # gpt-4o/mini/o3 系列
        return 128000
    if "gpt-4" in m:              # 经典 gpt-4/0613 等
        return 8192
    if "gpt-3.5" in m:
        return 4096
    return int(os.getenv("OPENAI_CONTEXT_LIMIT", "8192"))

def est_tokens_from_messages(msgs: list[dict]) -> int:
    # 粗估：字符数 / 3.6 ≈ tokens
    s = "".join((m.get("content") or "") for m in msgs)
    return max(1, int(len(s) / 3.6))

def shrink_metrics(metric_summary: dict, keep_points: int) -> dict:
    out = {}
    for k, v in (metric_summary or {}).items():
        series = list(v.get("series") or [])
        if len(series) > keep_points:
            series = series[-keep_points:]
        out[k] = {**v, "series": series}
    return out

def shrink_policy(policy_ctx: list, limit: int, max_chars: int = 240) -> list:
    out = []
    for it in (policy_ctx or [])[:limit]:
        s = (it.get("summary") or "")[:max_chars]
        out.append({**it, "summary": s})
    return out
if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
    raise RuntimeError("缺少 SUPABASE_URL 或 SUPABASE_SERVICE_ROLE_KEY")

sb: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)
llm = OpenAI(api_key=OPENAI_API_KEY or None, base_url=OPENAI_BASE_URL)
def ensure_bucket(bucket: str):
    try:
        buckets = sb.storage.list_buckets() or []
        names = [ (b.get("name") if isinstance(b, dict) else getattr(b, "name", None)) for b in buckets ]
        if bucket not in names:
            sb.storage.create_bucket(bucket, {"public": True, "file_size_limit": 104857600})
            logger.info("Created storage bucket: %s", bucket)
    except Exception as e:
        logger.warning("ensure_bucket failed: %s", e)

ensure_bucket(REPORTS_BUCKET)


# -------------------- FastAPI --------------------
app = FastAPI(title="Report Agent (Locked Prompt)", version="1.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"],
)

# -------------------- DTO --------------------
class Quarter(BaseModel):
    year: int
    quarter: str  # 'Q1'|'Q2'|'Q3'|'Q4'

class Params(BaseModel):
    company_name: str
    start: Quarter
    end: Quarter

class GeneratePayload(BaseModel):
    reportType: str = "annual_financial"
    language: str = "zh"
    specialRequirements: Optional[str] = None       # “特殊要求（可选）”
    templateId: Optional[str] = None                # 从DB读取
    templateData: Optional[Dict[str, Any]] = None   # 或者直接传模板JSON
    parameters: Params

# -------------------- 安全 --------------------
def auth_check(authorization: Optional[str] = Header(None)):
    if not REPORT_AGENT_TOKEN:
        return True
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(401, "Missing token")
    if authorization.split(" ", 1)[1] != REPORT_AGENT_TOKEN:
        raise HTTPException(403, "Invalid token")
    return True

# -------------------- 小工具 --------------------
def q_to_int(q: str) -> int:
    return {"Q1":1, "Q2":2, "Q3":3, "Q4":4}.get(q.upper(), 1)
def _norm_quarter(v) -> str:
    """把 1/2/3/4 或 '1'/'Q1' 统一成 'Q1'~'Q4'"""
    s = str(v).strip().upper()
    if s.startswith("Q"):
        try:
            return f"Q{int(s[1:])}"
        except Exception:
            return "Q1"
    try:
        n = int(float(s))
    except Exception:
        n = 1
    if n < 1 or n > 4:
        n = 1
    return f"Q{n}"
def period_key(y: int, q: str) -> int:
    return y*10 + q_to_int(q)

# -------------------- 模板（结构来自 DB 或请求体；提示词不在模板里） --------------------
DEFAULT_TEMPLATE = {
    "name": "年度财务报告（默认）",
    "sections": [
        {"title":"概述","hint":"年度概览与关键结论"},
        {"title":"第一章 经营与盈利能力","hint":"核心指标、同比/环比、杜邦拆解","requireCharts":True},
        {"title":"第二章 营运效率与现金流","hint":"周转率、现金流三大活动","requireCharts":True},
        {"title":"第三章 财务结构与风险点","hint":"资产负债率、偿债能力、主要风险"},
        {"title":"总结与行动建议","hint":"面向管理层的重点建议"}
    ],
    "required_metrics": [
        "营业收入","净利润","ROE","ROA",
        "毛利率","期间费用率",
        "经营活动现金流净额","投资活动现金流净额","筹资活动现金流净额",
        "总资产周转率","存货周转率","应收账款周转率","固定资产周转率",
        "资产负债率","流动比率","速动比率","现金比率"
    ],
    "variables":["company","start_year","start_quarter","end_year","end_quarter","language"]
}

def fetch_template(template_id: Optional[str], template_data: Optional[Dict[str, Any]]) -> Dict:
    if template_data:
        return template_data
    if not template_id:
        return DEFAULT_TEMPLATE
    try:
        res = sb.table("report_templates").select("*").eq("id", template_id).single().execute()
    except Exception as e:
        raise HTTPException(400, f"模板读取失败: {e}")
    row = (getattr(res, "data", None) or {})
    td = row.get("template_data") or {}
    return td if td else {**DEFAULT_TEMPLATE, "name": row.get("name","报告")}
def infer_required_metrics(template: Dict[str, Any]) -> List[str]:
    if isinstance(template.get("required_metrics"), list):
        return [str(x) for x in template["required_metrics"]]
    return DEFAULT_TEMPLATE["required_metrics"]

# -------------------- 数据读取与汇总 --------------------
def fetch_financial_metrics(company: str, start: Quarter, end: Quarter, metrics: List[str]) -> pd.DataFrame:
    years = list(range(start.year, end.year + 1))
    res = (
        sb.table("financial_metrics")
        .select("company_name, year, quarter, metric_name, metric_value")
        .eq("company_name", company)
        .in_("year", years)
        .in_("metric_name", metrics)
        .limit(50000)
        .execute()
    )

    df = pd.DataFrame(getattr(res, "data", []) or [])
    if df.empty:
        return df

    # 统一类型/格式
    df["year"] = df["year"].astype(int)

    def _norm_quarter(v) -> str:
        s = str(v).strip().upper()
        if s.startswith("Q"):
            try:
                return f"Q{int(s[1:])}"
            except Exception:
                return "Q1"
        try:
            n = int(float(s))
        except Exception:
            n = 1
        if n < 1 or n > 4:
            n = 1
        return f"Q{n}"

    df["quarter"] = df["quarter"].apply(_norm_quarter)

    def period_key(y: int, q: str) -> int:
        mp = {"Q1": 1, "Q2": 2, "Q3": 3, "Q4": 4}
        return int(y) * 10 + mp.get(q.upper(), 1)

    df["pkey"] = df.apply(lambda r: period_key(int(r["year"]), r["quarter"]), axis=1)

    p_start = period_key(start.year, start.quarter)
    p_end   = period_key(end.year, end.quarter)
    df = df[(df["pkey"] >= p_start) & (df["pkey"] <= p_end)].copy()

    df.sort_values(by=["metric_name", "year", "quarter"], inplace=True)
    return df


def summarize_timeseries(df: pd.DataFrame) -> Dict[str, Any]:
    """
    输入：company_name, year, quarter, metric_name, metric_value, pkey
    输出：按 metric_name 的时序与基础统计
    """
    out: Dict[str, Any] = {}
    if df.empty:
        return out

    for name, grp in df.groupby("metric_name"):
        grp = grp.sort_values(["year", "pkey"])
        series = [
            {
                "period": f"{int(r.year)}{r.quarter}",
                "year": int(r.year),
                "quarter": r.quarter,
                "value": r.metric_value,
            }
            for r in grp.itertuples(index=False)
        ]
        last = series[-1]["value"] if series else None
        prev = series[-2]["value"] if len(series) > 1 else None
        yoy  = None
        if len(series) >= 5:
            base = series[-5]["value"]
            yoy = None if base in (None, 0) else (last - base) / base
        qoq = None if prev in (None, 0) else (last - prev) / prev
        out[name] = {"series": series, "latest": last, "qoq": qoq, "yoy": yoy}
    return out

# -------------------- 政策上下文 --------------------
def fetch_policy_from_table(limit: int = 8) -> List[Dict[str, Any]]:
    try:
        res = sb.table("policy_news").select("*").order("created_at", desc=True).limit(limit).execute()
        rows = getattr(res, "data", None) or []
    except Exception:
        rows = []
    items = []
    for r in rows:
        items.append({
            "title": r.get("title") or r.get("headline") or "",
            "url": r.get("url") or r.get("source") or "",
            "summary": r.get("summary") or r.get("content") or ""
        })
    return items

def bing_search(query: str, count: int = 5) -> List[Dict[str, str]]:
    headers = {"Ocp-Apim-Subscription-Key": BING_SUBSCRIPTION_KEY}
    params = {"q": query, "mkt":"zh-CN", "count": count, "freshness":"Year"}
    r = requests.get(BING_ENDPOINT, headers=headers, params=params, timeout=20)
    r.raise_for_status()
    js = r.json()
    out = []
    for v in (js.get("webPages") or {}).get("value", []):
        out.append({"title": v.get("name",""), "url": v.get("url",""), "summary": v.get("snippet","")})
    return out

def google_cse_search(query: str, count: int = 5) -> List[Dict[str,str]]:
    url = "https://www.googleapis.com/customsearch/v1"
    params = {"key": GOOGLE_API_KEY, "cx": GOOGLE_CSE_ID, "q": query, "num": min(count,10)}
    r = requests.get(url, params=params, timeout=20)
    r.raise_for_status()
    js = r.json()
    out = []
    for it in js.get("items", []):
        out.append({"title": it.get("title",""), "url": it.get("link",""), "summary": it.get("snippet","")})
    return out

def collect_policy_context(company: str, start: Quarter, end: Quarter, extra_hint: Optional[str]) -> List[Dict[str,str]]:
    ctx = fetch_policy_from_table(limit=8)
    query = f"{company} 行业 政策 影响 {start.year}-{end.year}"
    external = []
    try:
        if BING_SUBSCRIPTION_KEY:
            external = bing_search(query, count=5)
        elif GOOGLE_API_KEY and GOOGLE_CSE_ID:
            external = google_cse_search(query, count=5)
    except Exception:
        external = []
    return ctx + external

# -------------------- 导出 & 存储 --------------------
def export_docx(md_text: str) -> bytes:
    doc = Document()
    for line in md_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), 0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), 1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2)
        elif line.startswith("> "):
            p = doc.add_paragraph(); run = p.add_run(line[2:].strip()); run.italic = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

def export_pdf(md_text: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x, y = 40, height - 40
    for raw in md_text.splitlines():
        line = raw.replace("\t","    ")
        if y < 60:
            c.showPage(); y = height - 40
        c.drawString(x, y, line[:120])
        y -= 16
    c.save(); return buf.getvalue()
ECHARTS_BLOCK_RE = re.compile(r"```echarts\s*([\s\S]*?)```", re.MULTILINE)

def _points_to_option(obj: dict):
    """
    把老格式 {type,title,series:[{name,points:[{x,y}]}]} 转成标准 ECharts option
    """
    try:
        series = obj.get("series") or []
        if not series or "points" not in series[0]:
            return None
        pts = series[0].get("points") or []
        xs = [str(p.get("x")) for p in pts]
        ys = [p.get("y") for p in pts]
        return {
            "title": {"text": obj.get("title") or ""},
            "tooltip": {"trigger": "axis"},
            "xAxis": {"type": "category", "data": xs},
            "yAxis": {"type": "value"},
            "series": [{
                "name": series[0].get("name") or "series",
                "type": obj.get("type") or "line",
                "data": ys
            }]
        }
    except Exception:
        return None

def normalize_echarts_blocks(md_text: str) -> str:
    """
    扫描 ```echarts ...```，若是旧的 points 结构则转为标准 option；
    同时强制 legend/tooltip/grid/yAxis.scale，避免渲染时被遮挡或顶边。
    """
    def _sub(m):
        raw = m.group(1).strip()
        try:
            obj = json.loads(raw)
            opt = _points_to_option(obj) or obj
            if isinstance(opt, dict):
                # legend/tooltip
                opt.setdefault("legend", {})
                if opt["legend"].get("show") is None:
                    opt["legend"]["show"] = True
                opt["legend"].setdefault("top", 6)
                opt["legend"].setdefault("left", "center")
                opt.setdefault("tooltip", {"trigger": "axis"})
                # 网格留白
                g = opt.setdefault("grid", {})
                g.setdefault("top", 48)
                g.setdefault("left", 56)
                g.setdefault("right", 32)
                g.setdefault("bottom", 48)
                g.setdefault("containLabel", True)
                # y 轴可缩放
                if isinstance(opt.get("yAxis"), list):
                    for y in opt["yAxis"]:
                        if isinstance(y, dict):
                            y.setdefault("type", "value")
                            y.setdefault("scale", True)
                elif isinstance(opt.get("yAxis"), dict):
                    opt["yAxis"].setdefault("type", "value")
                    opt["yAxis"].setdefault("scale", True)
                else:
                    opt["yAxis"] = {"type": "value", "scale": True}
                # series.name 补齐
                series = opt.get("series") if isinstance(opt.get("series"), list) else []
                names = []
                for i, s in enumerate(series):
                    if isinstance(s, dict) and not s.get("name"):
                        s["name"] = f"系列{i+1}"
                    if isinstance(s, dict):
                        names.append(s.get("name") or f"系列{i+1}")
                if names and not opt["legend"].get("data"):
                    opt["legend"]["data"] = names
            return "```echarts\n" + json.dumps(opt, ensure_ascii=False) + "\n```"
        except Exception:
            return m.group(0)
    return ECHARTS_BLOCK_RE.sub(_sub, md_text)



# --- 改后（整段替换） ---
def upload_bytes_to_storage(path: str, content: bytes, content_type="application/octet-stream") -> str:
    sb.storage.from_(REPORTS_BUCKET).upload(
        path,
        content,
        {"contentType": str(content_type), "upsert": "true"}
    )
    return sb.storage.from_(REPORTS_BUCKET).get_public_url(path)

# -------------------- 锁定提示词（前端无法覆盖） --------------------
STRICT_SYSTEM_PROMPT = (
    "你是资深企业财务分析师与报告撰写专家。你的任务：基于给定的‘时间范围、公司、模板结构、"
    "财务指标时间序列（含同比/环比）、政策上下文’生成**严谨、客观、条理清晰**的管理层阅读报告（Markdown）。\n"
    "必须遵守：\n"
    "1) 不得编造数据；仅使用我提供的数据与政策来源。未知处写“数据不足，建议补充”。\n"
    "2) 金额保留两位小数；比例以百分号表示（如 12.34%）。\n"
    "3) 避免情绪化措辞；给出可执行建议。\n"
    # 4) 图表使用如下代码块（ECharts 标准 option）
    "4) 图表使用如下代码块（ECharts 标准 option）：\n"
    "```echarts\n"
    "{\n"
    "  \"title\": {\"text\": \"...\"},\n"
    "  \"tooltip\": {\"trigger\": \"axis\"},\n"
    "  \"xAxis\": {\"type\": \"category\", \"data\": [\"2024Q1\",\"2024Q2\",\"2024Q3\",\"2024Q4\"]},\n"
    "  \"yAxis\": {\"type\": \"value\"},\n"
    "  \"series\": [{\"name\": \"指标名\", \"type\": \"line\", \"data\": [12.3, 15.1, 13.8, 16.4]}]\n"
    "}\n"
    "```\n"

    "5) 若使用政策上下文，段落末尾以 [n] 形式标注引用编号；结尾列出“参考来源”。\n"
)

STYLE_GUIDE = (
    "写作风格：\n"
    "- 每个章节以 3–6 条要点概括；要点下可有简短解释。\n"
    "- 结构顺序严格按照模板 sections；每章使用二级标题（##）。\n"
    "- 最后一章输出 2–5 条“管理层行动建议”。\n"
)

ACCURACY_RULES = (
    "准确性规则：\n"
    "- 同比=与上年同季度对比；环比=与上个季度对比。\n"
    "- 缺少基础值时，不计算同比/环比。\n"
    "- 单位来自数据表的 unit 字段；若缺失则不强行添加。\n"
)

# -------------------- 构建 LLM 消息（锁定） --------------------
def build_locked_messages(template: Dict[str, Any], company: str, start: Quarter, end: Quarter,
                          metric_summary: Dict[str, Any], policy_ctx: List[Dict[str,str]],
                          language: str, special: Optional[str],
                          keep_points: int = 24, policy_limit: int = 8) -> List[Dict[str,str]]:
    # 压缩时序 & 政策
    compact_metrics = shrink_metrics(metric_summary, keep_points=keep_points)
    framework = {"name": template.get("name","报告"), "sections": template.get("sections", [])}
    policy_slice = shrink_policy(policy_ctx, limit=policy_limit)

    locked_instruction = STRICT_SYSTEM_PROMPT + "\n" + STYLE_GUIDE + "\n" + ACCURACY_RULES
    user_payload = {
        "company": company,
        "period": {"start": start.dict(), "end": end.dict()},
        "language": language,
        "special_requirements": special or "",
        "report_framework": framework,
        "metrics": compact_metrics,
        "policy_context": policy_slice
    }
    return [
        {"role": "system", "content": locked_instruction},
        {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)}
    ]

def call_llm(messages):
    """
    只走 OpenAI 兼容接口：/v1/chat/completions
    基址从 OPENAI_BASE_URL 读取（必须包含 /v1）
    """
    base = (os.getenv("OPENAI_BASE_URL") or os.getenv("OPENAI_API_BASE") or "https://api.openai.com/v1").strip().rstrip("/")
    if not base.endswith("/v1"):
        base = base + "/v1"
    key  = os.getenv("OPENAI_API_KEY") or ""
    model = os.getenv("OPENAI_MODEL") or "gpt-4"
    
    # 避免被 Azure / 其他变量“劫持”
    os.environ.pop("AZURE_OPENAI_ENDPOINT", None)
    os.environ.pop("OPENAI_API_BASE", None)

    logger.info("LLM base=%s model=%s", base, model)

    try:
        client = OpenAI(base_url=base, api_key=key, timeout=60)
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.2
        )
        return resp.choices[0].message.content or ""
    except Exception as e:
        logger.error("LLM request failed. base=%s model=%s err=%s", base, model, e)
        # 让上层 generate() 的兜底捕到并返回 detail
        raise

# -------------------- 路由 --------------------
@app.get("/health")
def health():
    return {"ok": True, "time": dt.datetime.utcnow().isoformat(), "use_llm": USE_LLM}

@app.post("/report/generate")
def generate(payload: GeneratePayload, _=Depends(auth_check)):
    now = dt.datetime.utcnow().isoformat()
    try:
        p = payload.parameters

        # 1) 模板
        template = fetch_template(payload.templateId, payload.templateData)
        need_metrics = infer_required_metrics(template)

        # 2) 数据
        df = fetch_financial_metrics(p.company_name, p.start, p.end, need_metrics)
        metric_summary = summarize_timeseries(df)

        # 3) 政策
        policy_ctx = collect_policy_context(
            p.company_name, p.start, p.end, payload.specialRequirements
        )

        # 4) LLM（单独兜底：把 base/model 打到日志 & 返回 detail）
        try:
            llm_base = (
                os.getenv("OPENAI_BASE_URL")
                or os.getenv("OPENAI_API_BASE")
                or os.getenv("AZURE_OPENAI_ENDPOINT")
                or "https://api.openai.com/v1"
            )
            llm_model = os.getenv("OPENAI_MODEL") or os.getenv("AZURE_OPENAI_DEPLOYMENT") or "(unset)"
            messages = build_locked_messages(
                template, p.company_name, p.start, p.end,
                metric_summary, policy_ctx, payload.language, payload.specialRequirements
            )
            content_md = call_llm(messages)
            content_md = normalize_echarts_blocks(content_md)
        except Exception as e:
            logger.error("LLM request failed. base=%s model=%s err=%s\n%s",
                         llm_base, llm_model, e, traceback.format_exc())
            # 让前端能直接看到是 LLM 基址/路由问题
            raise HTTPException(status_code=500, detail=f"llm_failed: base={llm_base} model={llm_model} err={e}")

        # 5) 结果
        job_id = str(uuid.uuid4())
        result = {
            "job_id": job_id,
            "generated_at": now,
            "content_md": content_md,
            "metadata": {
                "reportType": payload.reportType,
                "language": payload.language,
                "sections": [s.get("title","章节") for s in template.get("sections",[])],
                "generatedAt": now,
                "dataRange": f"{p.start.year}{p.start.quarter}–{p.end.year}{p.end.quarter}",
                "aiGenerated": True
            },
            "logs": [
                {"step":"template_loaded","name": template.get("name")},
                {"step":"metrics_required","count": len(need_metrics)},
                {"step":"financial_rows","rows": int(df.shape[0]) if isinstance(df, pd.DataFrame) else 0},
                {"step":"policy_collected","count": len(policy_ctx)},
                {"step":"llm_done","chars": len(content_md)}
            ]
        }

        # 6) 导出（你已关闭就不动）
        if EXPORT_ENABLED:
            try:
                day = dt.datetime.utcnow().strftime("%Y%m%d")
                base = f"{day}/{job_id}"
                docx_bytes = export_docx(content_md)
                pdf_bytes  = export_pdf(content_md)
                docx_url = upload_bytes_to_storage(
                    f"{base}/report.docx", docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                pdf_url  = upload_bytes_to_storage(f"{base}/report.pdf",  pdf_bytes, "application/pdf")
                result.update({
                    "docx_url": docx_url,
                    "pdf_url": pdf_url,
                    "file_name": f"{p.company_name}_报告_{p.start.year}{p.start.quarter}-{p.end.year}{p.end.quarter}.pdf"
                })
                result["logs"].append({"step":"files_uploaded"})
            except Exception as e:
                logger.error("export failed: %s\n%s", e, traceback.format_exc())
                result["logs"].append({"step":"export_failed","error": str(e)})

        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error("generate failed: %s\n%s", e, traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"generate failed: {e}")



# -------------------- 新增：流式生成（SSE，带自动续写） --------------------
def _normalize_base_url() -> str:
    base = (os.getenv("OPENAI_BASE_URL") or os.getenv("OPENAI_API_BASE") or "https://api.openai.com/v1").strip().rstrip("/")
    if not base.endswith("/v1"):
        base += "/v1"
    return base

def _sse(event_type: str, data: dict | str) -> str:
    payload = data if isinstance(data, str) else json.dumps(data, ensure_ascii=False)
    return f"event: {event_type}\n" + f"data: {payload}\n\n"

@app.post("/report/stream")
def report_stream(payload: GeneratePayload, _=Depends(auth_check)):
    """
    SSE：event: progress|chunk|result|error|done
    """
    try:
        p = payload.parameters
    except Exception as e:
        raise HTTPException(400, f"bad payload: {e}")

    def _sse(event_type: str, data: dict | str) -> str:
        payload_ = data if isinstance(data, str) else json.dumps(data, ensure_ascii=False)
        return f"event: {event_type}\n" + f"data: {payload_}\n\n"

    def _normalize_base_url() -> str:
        base = (os.getenv("OPENAI_BASE_URL") or os.getenv("OPENAI_API_BASE") or "https://api.openai.com/v1").strip().rstrip("/")
        if not base.endswith("/v1"):
            base += "/v1"
        return base

    def _gen():
        now = dt.datetime.utcnow().isoformat()

        # 1) 模板
        try:
            yield _sse("progress", {"stage": "加载模板"})
            template = fetch_template(payload.templateId, payload.templateData)
            need_metrics = infer_required_metrics(template)
            section_titles = [s.get("title","") for s in template.get("sections",[])]
            yield _sse("progress", {"stage": "模板完成", "name": template.get("name","默认模板"), "metrics": len(need_metrics)})
        except Exception as e:
            yield _sse("error", {"message": f"模板读取失败: {e}"})
            yield _sse("done", {})
            return

        # 2) 数据
        try:
            yield _sse("progress", {"stage": "拉取财务数据"})
            df = fetch_financial_metrics(p.company_name, p.start, p.end, need_metrics)
            metric_summary = summarize_timeseries(df)
            rows = int(df.shape[0]) if isinstance(df, pd.DataFrame) else 0
            yield _sse("progress", {"stage": "数据就绪", "rows": rows, "metrics_found": len(metric_summary)})
        except Exception as e:
            yield _sse("error", {"message": f"数据查询失败: {e}"})
            yield _sse("done", {})
            return

        # 3) 政策
        try:
            yield _sse("progress", {"stage": "收集政策上下文"})
            policy_ctx = collect_policy_context(p.company_name, p.start, p.end, payload.specialRequirements)
            yield _sse("progress", {"stage": "政策上下文完成", "count": len(policy_ctx)})
        except Exception as e:
            policy_ctx = []
            yield _sse("progress", {"stage": "政策上下文跳过", "error": str(e)})

        # 4) 组装消息 + 预算
        try:
            base = _normalize_base_url()
            key  = os.getenv("OPENAI_API_KEY") or ""
            model = os.getenv("OPENAI_MODEL") or "gpt-4"
            ctx_limit = guess_context_limit(model)

            keep_points = 24
            policy_limit = 8

            def build_and_budget():
                msgs = build_locked_messages(
                    template, p.company_name, p.start, p.end,
                    metric_summary, policy_ctx, payload.language, payload.specialRequirements,
                    keep_points=keep_points, policy_limit=policy_limit
                )
                prompt_tokens = est_tokens_from_messages(msgs)
                reserve = 512
                max_out = max(512, min(3072, ctx_limit - prompt_tokens - reserve))
                return msgs, prompt_tokens, max_out

            messages, prompt_tokens, max_out = build_and_budget()
            tighten_round = 0
            while max_out < 900 and tighten_round < 2:
                tighten_round += 1
                keep_points = 12 if keep_points > 12 else 8
                policy_limit = 4 if policy_limit > 4 else 2
                messages, prompt_tokens, max_out = build_and_budget()

            yield _sse("progress", {
                "stage": "调用模型（流式）",
                "ctx_limit": ctx_limit, "prompt_tokens": prompt_tokens, "max_tokens": max_out,
                "keep_points": keep_points, "policy_limit": policy_limit
            })
        except Exception as e:
            yield _sse("error", {"message": f"消息构建失败: {e}"})
            yield _sse("done", {})
            return

        # 5) 调用 LLM（流式 + 自动续写）
        content_md = ""
        try:
            client = OpenAI(
                base_url=base,
                api_key=key,
                timeout=httpx.Timeout(connect=10.0, read=300.0, write=30.0, pool=10.0),
                max_retries=2,
            )

            def stream_once(msgs, round_no: int, max_tokens: int):
                nonlocal content_md
                finish_reason = None
                stream = client.chat.completions.create(
                    model=model,
                    messages=msgs,
                    temperature=0.2,
                    stream=True,
                    max_tokens=max_tokens,
                )
                for chunk in stream:
                    if not chunk.choices:
                        continue
                    choice = chunk.choices[0]
                    delta = choice.delta
                    if getattr(choice, "finish_reason", None):
                        finish_reason = choice.finish_reason
                    if delta and getattr(delta, "content", None):
                        txt = delta.content
                        content_md += txt
                        yield _sse("chunk", {"text": txt})
                    yield _sse("heartbeat", "1")
                content_md = normalize_echarts_blocks(content_md)
                return finish_reason

            # 章节完成度判断
            section_titles = [s.get("title","") for s in template.get("sections", [])]
            def count_done_sections(text: str) -> int:
                return sum(1 for t in section_titles if t and (f"## {t}" in text))

            fr = yield from stream_once(messages, 1, max_out)
            done = count_done_sections(content_md)
            need_continue = (fr == "length") or (done < len(section_titles))

            rounds = 0
            while need_continue and rounds < 3:
                rounds += 1
                missing = [t for t in section_titles if t and (f"## {t}" not in content_md)]
                hint = "、".join(missing) if missing else "剩余章节"
                yield _sse("progress", {"stage": f"继续生成（第{rounds}轮）", "missing": missing})

                continue_messages = messages + [
                    {"role": "assistant", "content": content_md},
                    {"role": "user", "content":
                        f"请从上次中断处继续写作，补齐未完成章节（{hint}）。"
                        "不要重复已写内容，延续编号与格式，直到所有章节输出完毕。"}
                ]
                fr = yield from stream_once(continue_messages, rounds + 1, max_out)
                done = count_done_sections(content_md)
                need_continue = (fr == "length") or (done < len(section_titles))

            yield _sse("progress", {"stage": "生成完成"})
        except Exception as e:
            logger.error("LLM stream failed: %s\n%s", e, traceback.format_exc())
            yield _sse("error", {"message": f"llm_failed: base={base} model={model} err={e}"})
            yield _sse("done", {})
            return

        # 6) 导出与结果
        try:
            job_id = str(uuid.uuid4())
            result = {
                "job_id": job_id,
                "generated_at": now,
                "content_md": content_md,
                "metadata": {
                    "reportType": payload.reportType,
                    "language": payload.language,
                    "sections": section_titles,
                    "generatedAt": now,
                    "dataRange": f"{p.start.year}{p.start.quarter}–{p.end.year}{p.end.quarter}",
                    "aiGenerated": True
                }
            }

            if EXPORT_ENABLED and content_md.strip():
                try:
                    day = dt.datetime.utcnow().strftime("%Y%m%d")
                    base_path = f"{day}/{job_id}"
                    docx_bytes = export_docx(content_md)
                    pdf_bytes  = export_pdf(content_md)
                    docx_url = upload_bytes_to_storage(
                        f"{base_path}/report.docx", docx_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    pdf_url  = upload_bytes_to_storage(f"{base_path}/report.pdf", pdf_bytes, "application/pdf")
                    result.update({
                        "docx_url": docx_url,
                        "pdf_url": pdf_url,
                        "file_name": f"{p.company_name}_报告_{p.start.year}{p.start.quarter}-{p.end.year}{p.end.quarter}.pdf"
                    })
                except Exception as e:
                    result["export_error"] = str(e)

            yield _sse("result", result)
        except Exception as e:
            yield _sse("error", {"message": f"result_failed: {e}"})
        finally:
            yield _sse("done", {})

    return StreamingResponse(
        _gen(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"}
    )

